# main.py
# import json
import logging
import os
import uuid
from datetime import datetime
from io import BytesIO

import requests
import streamlit as st
from bs4 import BeautifulSoup
from docx import Document
from dotenv import load_dotenv
from streamlit_lottie import st_lottie

# Logging configuration
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()


def load_lottie(url: str):
    """Load animation from Lottie Files."""
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()


class ArticlesTranslator:
    """Class for translating articles using Azure OpenAI."""

    def __init__(self):
        self.api_key = os.getenv("AZURE_OPENAI_KEY")
        self.endpoint = os.getenv("AZURE_ENDPOINT")

    def extract_text(self, url: str) -> str:
        try:
            response = requests.get(url)
            response.raise_for_status()

            soup = BeautifulSoup(response.text, "html.parser")
            for script in soup(["script", "style"]):
                script.decompose()

            return soup.get_text(" ", strip=True)

        except requests.RequestException as e:
            logger.error(f"Failed to fetch URL: {e}")
            raise

    def translate_text(self, text: str, target_language: str) -> str:
        headers = {
            "Content-Type": "application/json",
            "api-key": self.api_key,
        }

        payload = {
            "messages": [
                {
                    "role": "system",
                    "content": [
                        {"type": "text", "text": "You act as a text translator"}
                    ],
                },
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": f"translate: {text} to {target_language} language and respond only with the translation in markdown format",
                        }
                    ],
                },
            ],
            "temperature": 0.9,
            "top_p": 0.95,
            "max_tokens": 900,
        }

        try:
            response = requests.post(
                self.endpoint, headers=headers, json=payload)
            response.raise_for_status()
            return response.json()["choices"][0]["message"]["content"]

        except requests.RequestException as e:
            logger.error(f"API request failed: {e}")
            raise


class DocumentsTranslator:
    """Class for translating Word documents using Azure Translator."""

    def __init__(self):
        self.api_key = os.getenv("TRANSLATOR_API_KEY")
        self.endpoint = os.getenv("TRANSLATOR_ENDPOINT")
        self.location = os.getenv("TRANSLATOR_LOCATION")

    def translate_text(self, text, source_language, target_language):
        path = "/translate"
        constructed_url = self.endpoint + path

        params = {"api-version": "3.0",
                  "from": source_language, "to": [target_language]}

        headers = {
            "Ocp-Apim-Subscription-Key": self.api_key,
            "Ocp-Apim-Subscription-Region": self.location,
            "Content-type": "application/json",
            "X-ClientTraceId": str(uuid.uuid4()),
        }

        body = [{"text": text}]
        response = requests.post(
            constructed_url, params=params, headers=headers, json=body
        )

        if response.status_code == 200:
            return response.json()[0]["translations"][0]["text"]
        else:
            raise Exception(f"Translation error: {response.status_code}")


def home_page():
    """Application home page."""
    st.title("🌍 Multilingual Translator")

    # Load animation
    lottie_translate = load_lottie(
        "https://lottie.host/e61f4b1a-23b5-4f25-b0f8-ec1f3569f717/1EfdaRqh4y.json"
    )
    st_lottie(lottie_translate, height=300)

    st.markdown(
        """
    ## Welcome to the Multilingual Translator!
    # Our platform offers two powerful translation tools:
    ### 📰 Article Translator
    Translate web articles while maintaining markdown formatting and original structure.
    ### 📄 Document Translator
    Translate Word documents with support for multiple languages and direct download.
    ### Technologies Used
    #### Article Translator:
    - Azure OpenAI GPT-4o mini
    - BeautifulSoup4
    - Requests
    - Streamlit
    #### Document Translator:
    - Azure Translator API
    - python-docx
    - Streamlit
    - UUID
    ### About the Project
    This project was developed to facilitate content translation in different formats,
    offering an intuitive interface and professional results.
    """
    )


def article_translator_page():
    """Article translator page."""
    st.title("📰 Article Translator")

    # Load animation
    lottie_article = load_lottie(
        "https://lottie.host/cdc7f167-7c8a-4b67-9a5f-3db936a8cb8d/IB6FNEMBAR.json"
    )
    st_lottie(lottie_article, height=200)

    st.markdown(
        """
    ### How does it work?
    1. Paste the URL of the article you want to translate
    2. Select the target language
    3. Click translate
    4. Download the result in markdown format
    """
    )

    translator = ArticlesTranslator()

    url = st.text_input("🔗 Article URL")

    languages = {
        "Portuguese": "português",
        "English": "english",
        "Spanish": "español",
        "French": "français",
        "German": "deutsch",
        "Italian": "italiano",
    }

    target_language = st.selectbox("🌍 Target Language", list(languages.keys()))

    if st.button("🔄 Translate Article"):
        try:
            with st.spinner("Processing article..."):
                text = translator.extract_text(url)
                translation = translator.translate_text(
                    text, languages[target_language])

                st.markdown("### Translation Result")
                st.markdown(translation)

                # Download button
                st.download_button(
                    label="📥 Download Translation",
                    data=translation.encode("utf-8"),
                    file_name=f"translation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                    mime="text/markdown",
                )

        except Exception as e:
            st.error(f"An error occurred: {str(e)}")

    st.markdown("---")


def document_translator_page():
    """Word document translator page."""
    st.title("📄 Word Document Translator")

    # Load animation
    lottie_doc = load_lottie(
        "https://lottie.host/a4fdcc12-f826-4986-a561-6f06a3bfb1e1/kFJXH2m2jf.json"
    )
    st_lottie(lottie_doc, height=200)

    st.markdown(
        """
    ### How does it work?
    1. Upload your Word document
    2. Select source and target languages
    3. Click translate
    4. Download the translated document
    """
    )

    translator = DocumentsTranslator()

    file = st.file_uploader("📎 Upload Word file", type=["docx"])

    languages = {
        "English": "en",
        "French": "fr",
        "Spanish": "es",
        "German": "de",
        "Portuguese": "pt",
        "Italian": "it",
    }

    col1, col2 = st.columns(2)
    with col1:
        source_language = st.selectbox(
            "🔤 Source language", list(languages.keys()))
    with col2:
        target_language = st.selectbox(
            "🔤 Target language", list(languages.keys()), index=4
        )

    if st.button("🔄 Translate"):
        if file is not None:
            try:
                with st.spinner("Translating document..."):
                    # Load document
                    doc = Document(file)
                    original_text = "\n".join(
                        [paragraph.text for paragraph in doc.paragraphs]
                    )

                    # Translate
                    translation = translator.translate_text(
                        original_text, languages[source_language], languages[target_language]
                    )
                    st.markdown("### Original Text")
                    st.text_area("original text", original_text, height=200)
                    st.markdown("### Translation Result")
                    st.text_area("translation", translation, height=200)

                    # Create new document
                    translated_doc = Document()
                    for line in translation.split("\n"):
                        translated_doc.add_paragraph(line)

                    # Prepare for download
                    buffer = BytesIO()
                    translated_doc.save(buffer)
                    buffer.seek(0)

                    # Download button
                    st.download_button(
                        label="📥 Download Translated Document",
                        data=buffer,
                        file_name=f"translation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

            except Exception as e:
                st.error(f"An error occurred: {str(e)}")
        else:
            st.error("Please upload a Word file.")

    st.markdown("---")


def main():
    """Main application function."""
    st.set_page_config(page_title="Multilingual Translator", layout="wide")

    # Sidebar menu
    st.sidebar.title("📚 Menu")
    pages = {
        "🏠 Home Page": home_page,
        "📰 Article Translator": article_translator_page,
        "📄 Document Translator": document_translator_page,
    }

    choice = st.sidebar.radio("Navigation", list(pages.keys()))

    # Menu footer
    st.sidebar.markdown("---")
    st.sidebar.markdown(

    )

    # Render selected page
    pages[choice]()


if __name__ == "__main__":
    main()
